from __future__ import annotations

import uuid
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from dharmiq.db.models.uploads import ProcessingStage, UserUploadChunk
from dharmiq.db.models.users import User
from dharmiq.db.session import get_session_factory
from dharmiq.ingestion.parser import PageText
from dharmiq.core.errors import IngestionError
from dharmiq.ingestion.upload_pipeline import (
    create_user_upload,
    process_user_upload,
    process_user_upload_safe,
)
from dharmiq.llm.embeddings import EmbeddingBackend


class _FixedEmbeddingBackend(EmbeddingBackend):
    def __init__(self, *, dimensions: int = 384) -> None:
        self._dimensions = dimensions

    @property
    def dimensions(self) -> int:
        return self._dimensions

    async def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [[0.1] * self._dimensions for _ in texts]


class _StubParser:
    def extract_pages(self, file_path: Path) -> list[PageText]:
        return [
            PageText(
                page_number=1,
                text="Section 1. Notice period.\nThe employee is entitled to thirty days notice.",
            )
        ]


def _pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


@pytest.fixture(autouse=True)
async def _clean_uploads() -> None:
    factory = get_session_factory()
    async with factory() as db:
        await db.execute(text("DELETE FROM chat_session_uploads"))
        await db.execute(text("DELETE FROM user_upload_chunks"))
        await db.execute(text("DELETE FROM user_uploads"))
        await db.commit()
    yield


async def _create_user(db: AsyncSession) -> User:
    user = User(
        email=f"upload-{uuid.uuid4()}@example.com",
        hashed_password="test",
        is_active=True,
        is_superuser=False,
        is_verified=True,
    )
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return user


class _FailingParser:
    def extract_pages(self, file_path: Path) -> list[PageText]:
        raise ValueError("corrupt pdf")


@pytest.mark.asyncio
async def test_upload_happy_path_stages(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))

    from dharmiq.config.settings import get_settings

    get_settings.cache_clear()
    settings = get_settings()
    content = _pdf_bytes()
    factory: async_sessionmaker[AsyncSession] = get_session_factory()
    parser = _StubParser()
    backend = _FixedEmbeddingBackend()

    async with factory() as db:
        user = await _create_user(db)
        upload = await create_user_upload(
            db,
            user_id=user.id,
            filename="contract.pdf",
            mime_type="application/pdf",
            content=content,
            settings=settings,
        )
        assert upload.processing_stage == ProcessingStage.UPLOADED.value

        chunk_count = await process_user_upload(
            db,
            upload.id,
            settings=settings,
            pdf_parser=parser,
            embedding_backend=backend,
        )
        await db.refresh(upload)

        assert upload.processing_stage == ProcessingStage.READY.value
        assert chunk_count > 0
        assert upload.chunk_count == chunk_count
        assert upload.processing_error is None
        child_chunks = (
            await db.execute(
                select(UserUploadChunk).where(
                    UserUploadChunk.upload_id == upload.id,
                    UserUploadChunk.parent_chunk_id.is_not(None),
                )
            )
        ).scalars().all()
        assert upload.chunk_count == len(child_chunks)


@pytest.mark.asyncio
async def test_upload_parse_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))

    from dharmiq.config.settings import get_settings

    get_settings.cache_clear()
    settings = get_settings()
    content = _pdf_bytes()
    factory: async_sessionmaker[AsyncSession] = get_session_factory()

    async with factory() as db:
        user = await _create_user(db)
        upload = await create_user_upload(
            db,
            user_id=user.id,
            filename="contract.pdf",
            mime_type="application/pdf",
            content=content,
            settings=settings,
        )

        with pytest.raises(IngestionError):
            await process_user_upload_safe(
                db,
                upload.id,
                settings=settings,
                pdf_parser=_FailingParser(),
            )

        await db.refresh(upload)
        assert upload.processing_stage == ProcessingStage.FAILED.value
        assert upload.processing_error is not None
        assert "Upload parsing failed" in upload.processing_error


@pytest.mark.asyncio
async def test_process_user_upload_is_idempotent(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))

    from dharmiq.config.settings import get_settings

    get_settings.cache_clear()
    settings = get_settings()

    content = _pdf_bytes()

    factory: async_sessionmaker[AsyncSession] = get_session_factory()
    parser = _StubParser()
    backend = _FixedEmbeddingBackend()

    async with factory() as db:
        user = await _create_user(db)
        upload = await create_user_upload(
            db,
            user_id=user.id,
            filename="contract.pdf",
            mime_type="application/pdf",
            content=content,
            settings=settings,
        )

        first = await process_user_upload(
            db,
            upload.id,
            settings=settings,
            pdf_parser=parser,
            embedding_backend=backend,
        )
        second = await process_user_upload(
            db,
            upload.id,
            settings=settings,
            pdf_parser=parser,
            embedding_backend=backend,
        )

        assert first == second
        assert first >= 1

        all_chunks = (
            await db.execute(
                select(UserUploadChunk).where(UserUploadChunk.upload_id == upload.id)
            )
        ).scalars().all()
        leaf_chunks = [chunk for chunk in all_chunks if chunk.parent_chunk_id is not None]
        assert len(leaf_chunks) == first


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Termination", level=1)
    document.add_paragraph("The employee must provide thirty days written notice.")
    document.add_heading("Benefits", level=1)
    document.add_paragraph("Health insurance continues through the notice period.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _markdown_bytes() -> bytes:
    content = """# Termination

The employee must provide thirty days written notice.

# Benefits

Health insurance continues through the notice period.
"""
    return content.encode("utf-8")


@pytest.mark.asyncio
async def test_docx_upload_parsed(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))

    from dharmiq.config.settings import get_settings

    get_settings.cache_clear()
    settings = get_settings()
    factory: async_sessionmaker[AsyncSession] = get_session_factory()
    backend = _FixedEmbeddingBackend()

    async with factory() as db:
        user = await _create_user(db)
        upload = await create_user_upload(
            db,
            user_id=user.id,
            filename="contract.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=_docx_bytes(),
            settings=settings,
        )
        chunk_count = await process_user_upload(
            db,
            upload.id,
            settings=settings,
            embedding_backend=backend,
        )
        assert chunk_count >= 2

        chunks = (
            await db.execute(
                select(UserUploadChunk).where(UserUploadChunk.upload_id == upload.id)
            )
        ).scalars().all()
        leaf_chunks = [chunk for chunk in chunks if chunk.parent_chunk_id is not None]
        assert len(leaf_chunks) == chunk_count
        labels = {
            chunk.chunk_metadata.get("section_label")
            for chunk in chunks
            if chunk.chunk_metadata.get("section_label")
        }
        assert any("Termination" in (label or "") for label in labels)


@pytest.mark.asyncio
async def test_md_upload_parsed(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))

    from dharmiq.config.settings import get_settings

    get_settings.cache_clear()
    settings = get_settings()
    factory: async_sessionmaker[AsyncSession] = get_session_factory()
    backend = _FixedEmbeddingBackend()

    async with factory() as db:
        user = await _create_user(db)
        upload = await create_user_upload(
            db,
            user_id=user.id,
            filename="contract.md",
            mime_type="text/markdown",
            content=_markdown_bytes(),
            settings=settings,
        )
        chunk_count = await process_user_upload(
            db,
            upload.id,
            settings=settings,
            embedding_backend=backend,
        )
        assert chunk_count >= 2

        chunks = (
            await db.execute(
                select(UserUploadChunk).where(UserUploadChunk.upload_id == upload.id)
            )
        ).scalars().all()
        section_labels = [
            chunk.chunk_metadata.get("section_label")
            for chunk in chunks
            if chunk.chunk_metadata.get("section_label")
        ]
        assert any("Termination" in (label or "") for label in section_labels)
        assert any("Benefits" in (label or "") for label in section_labels)
