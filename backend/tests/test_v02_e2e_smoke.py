from __future__ import annotations

import asyncio
import json
import uuid
from datetime import UTC, datetime
from io import BytesIO
from unittest.mock import patch

import pytest
from docx import Document
from httpx import AsyncClient
from pgvector import Vector as PgVector
from sqlalchemy import func, select, text
from sqlalchemy.ext.asyncio import AsyncSession

from dharmiq.agents.runner import run_agent_graph_for_request
from dharmiq.config.settings import get_settings
from dharmiq.db.models.chats import ChatMessage, ChatRequestEvent, MessageRole
from dharmiq.db.models.documents import DocType, DocumentChunk, SourceDocument
from dharmiq.db.session import get_session_factory
from dharmiq.llm.embeddings import EmbeddingBackend
from dharmiq.ingestion.upload_pipeline import process_user_upload
from tests.litellm_helpers import mock_litellm_acompletion
from tests.rerank_helpers import mock_rerank
from tests.vector_helpers import unit_vector


class _StaticEmbeddingBackend(EmbeddingBackend):
    @property
    def dimensions(self) -> int:
        return 384

    async def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [unit_vector(index % 384) for index, _ in enumerate(texts)]


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Termination", level=1)
    document.add_paragraph("The employee must provide thirty days written notice.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _mock_full_pipeline_llm(monkeypatch: pytest.MonkeyPatch) -> None:
    clarifier = {
        "topic": "police_arrest",
        "needs_more_info": False,
        "followup_questions": [],
        "reason": "Enough detail",
    }
    rewriter = {"queries": ["Article 22 arrest rights", "employment contract notice period"]}
    answer = (
        "You have protections under Article 22 [1].\n\n"
        "> Article 22 protects against arbitrary arrest and detention.\n\n"
        "This is general legal information, not legal advice. Consult a qualified lawyer."
    )
    validator = {
        "must_regenerate": False,
        "issues": [],
        "regeneration_instructions": "",
        "final_warning": "Consult a qualified lawyer for your situation.",
    }
    mock_litellm_acompletion(
        monkeypatch,
        [
            json.dumps(clarifier),
            json.dumps(rewriter),
            answer,
            json.dumps(validator),
        ],
    )


async def _seed_corpus(db: AsyncSession) -> None:
    document = SourceDocument(
        source_id=f"e2e-{uuid.uuid4()}",
        title="Constitution of India (test)",
        doc_type=DocType.ACT,
        jurisdiction="central",
        content_hash="hash-e2e-smoke",
        file_path="/tmp/constitution.pdf",
        indexed_at=datetime.now(UTC),
    )
    db.add(document)
    await db.flush()
    db.add(
        DocumentChunk(
            document_id=document.id,
            chunk_index=0,
            text="Article 22 protects against arbitrary arrest and detention.",
            page_start=1,
            page_end=1,
            embedding=PgVector(unit_vector(0)),
        )
    )
    db.add(
        DocumentChunk(
            document_id=document.id,
            chunk_index=1,
            text="Article 22 also requires the grounds of arrest to be communicated.",
            page_start=1,
            page_end=1,
            embedding=PgVector(unit_vector(1)),
        )
    )
    await db.commit()


def _parse_sse_block(block: str) -> tuple[str | None, dict | None]:
    event_type = None
    data = None
    for line in block.strip().splitlines():
        if line.startswith("event: "):
            event_type = line.removeprefix("event: ")
        elif line.startswith("data: "):
            data = json.loads(line.removeprefix("data: "))
    return event_type, data


async def _read_sse_events(client: AsyncClient, url: str, headers: dict[str, str]) -> list[tuple[str, dict]]:
    events: list[tuple[str, dict]] = []
    async with client.stream("GET", url, headers=headers) as response:
        assert response.status_code == 200
        buffer = ""
        async for chunk in response.aiter_text():
            buffer += chunk
            while "\n\n" in buffer:
                part, buffer = buffer.split("\n\n", 1)
                if not part.strip():
                    continue
                event_type, data = _parse_sse_block(part)
                if event_type and data:
                    events.append((event_type, data))
    return events


async def _run_graph_inline(chat_request_id: uuid.UUID) -> None:
    factory = get_session_factory()
    async with factory() as db:
        with patch("dharmiq.llm.retrieval.get_embedding_backend", return_value=_StaticEmbeddingBackend()):
            await run_agent_graph_for_request(db, chat_request_id)


def _patch_inline_enqueue(monkeypatch: pytest.MonkeyPatch) -> list[asyncio.Task[None]]:
    tasks: list[asyncio.Task[None]] = []

    def enqueue(chat_request_id: uuid.UUID) -> None:
        tasks.append(asyncio.create_task(_run_graph_inline(chat_request_id)))

    monkeypatch.setattr("dharmiq.api.routes.chat.enqueue_agent_graph", enqueue)
    return tasks


def _patch_inline_upload_process(monkeypatch: pytest.MonkeyPatch) -> list[asyncio.Task[None]]:
    upload_tasks: list[asyncio.Task[None]] = []

    def send_task(name: str, args: list[str] | None = None, **_kwargs) -> None:
        del name

        async def process() -> None:
            upload_id = uuid.UUID(args[0])
            factory = get_session_factory()
            settings = get_settings()
            async with factory() as db:
                with patch(
                    "dharmiq.ingestion.upload_pipeline.get_embedding_backend",
                    return_value=_StaticEmbeddingBackend(),
                ):
                    await process_user_upload(db, upload_id, settings=settings)

        upload_tasks.append(asyncio.create_task(process()))

    monkeypatch.setattr("dharmiq.api.routes.uploads.celery_app.send_task", send_task)
    return upload_tasks


@pytest.fixture(autouse=True)
async def _clean_e2e_tables() -> None:
    factory = get_session_factory()
    async with factory() as db:
        await db.execute(text("DELETE FROM chat_request_events"))
        await db.execute(text("DELETE FROM chat_session_uploads"))
        await db.execute(text("DELETE FROM user_upload_chunks"))
        await db.execute(text("DELETE FROM user_uploads"))
        await db.execute(text("DELETE FROM chat_requests"))
        await db.execute(text("DELETE FROM chat_messages"))
        await db.execute(text("DELETE FROM chat_sessions"))
        await db.execute(text("DELETE FROM document_chunks"))
        await db.execute(text("DELETE FROM document_sections"))
        await db.execute(text("DELETE FROM source_documents"))
        await db.commit()
    yield


@pytest.fixture
def agent_graph_v2(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("DHARMIQ_AGENT_GRAPH_V2", "true")
    get_settings.cache_clear()


@pytest.mark.asyncio
@pytest.mark.timeout(120)
async def test_v02_e2e_smoke(
    client: AsyncClient,
    auth_headers: dict[str, str],
    agent_graph_v2: None,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path,
) -> None:
    """Full v0.2 path: register, session, DOCX upload, attach, message, SSE (mocked LLM)."""
    monkeypatch.setenv("DHARMIQ_ROOT", str(tmp_path))
    get_settings.cache_clear()

    _mock_full_pipeline_llm(monkeypatch)
    mock_rerank(monkeypatch)
    graph_tasks = _patch_inline_enqueue(monkeypatch)
    upload_tasks = _patch_inline_upload_process(monkeypatch)

    factory = get_session_factory()
    async with factory() as db:
        await _seed_corpus(db)

    session_resp = await client.post("/api/chat/sessions", json={"title": "E2E smoke"}, headers=auth_headers)
    assert session_resp.status_code == 201
    session_id = session_resp.json()["id"]

    upload_resp = await client.post(
        "/api/uploads",
        headers=auth_headers,
        files={
            "file": (
                "contract.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_resp.status_code == 201
    upload_id = upload_resp.json()["id"]

    if upload_tasks:
        await asyncio.gather(*upload_tasks)

    attach_resp = await client.post(
        f"/api/chat/sessions/{session_id}/attachments",
        headers=auth_headers,
        json={"upload_ids": [upload_id]},
    )
    assert attach_resp.status_code == 200
    assert len(attach_resp.json()) == 1

    post_resp = await client.post(
        f"/api/chat/sessions/{session_id}/messages",
        headers=auth_headers,
        json={"content": "What are my rights if police arrest me without a warrant?"},
    )
    assert post_resp.status_code == 202
    req_id = post_resp.json()["chat_request_id"]
    assert post_resp.json()["status"] == "pending"

    sse_task = asyncio.create_task(
        _read_sse_events(client, f"/api/chat/requests/{req_id}/stream", auth_headers)
    )
    await asyncio.gather(*graph_tasks)
    events = await sse_task

    done_events = [data for event_type, data in events if event_type == "done"]
    assert len(done_events) == 1
    assert done_events[0]["status"] == "completed"

    async with factory() as db:
        event_count = (
            await db.execute(
                select(func.count())
                .select_from(ChatRequestEvent)
                .where(ChatRequestEvent.chat_request_id == uuid.UUID(req_id))
            )
        ).scalar_one()
        assert event_count >= 5

        result = await db.execute(
            select(ChatMessage)
            .where(ChatMessage.message_metadata["chat_request_id"].astext == req_id)
            .where(ChatMessage.role == MessageRole.ASSISTANT)
        )
        assistant_msg = result.scalar_one()

    metadata = assistant_msg.message_metadata or {}
    assert metadata.get("citations")
    assert "[" in assistant_msg.content
    assert "legal advice" in assistant_msg.content.lower()

    citation_events = [data for event_type, data in events if event_type == "citation"]
    assert len(citation_events) >= 1


@pytest.mark.asyncio
@pytest.mark.timeout(60)
async def test_v05_export_delete_smoke(
    client: AsyncClient,
    unique_email: str,
) -> None:
    """v0.5 smoke: register, export account JSON, delete account, verify auth revoked."""
    password = "securepassword123"
    register = await client.post(
        "/api/auth/register",
        json={"email": unique_email, "password": password},
    )
    assert register.status_code == 201, register.text

    login = await client.post(
        "/api/auth/jwt/login",
        data={"username": unique_email, "password": password},
    )
    assert login.status_code == 200, login.text
    headers = {"Authorization": f"Bearer {login.json()['access_token']}"}

    session_resp = await client.post("/api/chat/sessions", json={"title": "Export smoke"}, headers=headers)
    assert session_resp.status_code == 201
    session_id = session_resp.json()["id"]

    message_resp = await client.post(
        f"/api/chat/sessions/{session_id}/messages",
        json={"role": "user", "content": "What are my fundamental rights?"},
        headers=headers,
    )
    assert message_resp.status_code == 201

    export_resp = await client.get("/api/account/export", headers=headers)
    assert export_resp.status_code == 200, export_resp.text
    assert export_resp.headers["content-type"].startswith("application/json")

    payload = export_resp.json()
    assert "exported_at" in payload
    assert payload["user"]["email"] == unique_email
    assert len(payload["sessions"]) == 1
    assert payload["sessions"][0]["id"] == session_id
    assert len(payload["messages"]) == 1
    assert payload["messages"][0]["session_id"] == session_id
    assert "uploads" in payload

    delete_resp = await client.request(
        "DELETE",
        "/api/account",
        headers=headers,
        json={"email": unique_email, "password": password},
    )
    assert delete_resp.status_code == 204, delete_resp.text

    export_after = await client.get("/api/account/export", headers=headers)
    assert export_after.status_code == 401
